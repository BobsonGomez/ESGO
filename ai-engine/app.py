from flask import Flask, request, jsonify
from google import genai
from google.genai import types
import PyPDF2
import docx
import io
import json
import os
import time

app = Flask(__name__)

# =========================================================
# 1. CONFIGURATION
# =========================================================

# PASTE YOUR API KEY HERE
os.environ["GOOGLE_API_KEY"] = "AIzaSyBL0rMvtS4FeogdVOtxKSXa76aaQ69vqNw"

client = genai.Client(api_key=os.environ["GOOGLE_API_KEY"])

# Gemini 1.5 Flash is best for large document analysis
MODEL_NAME = 'gemini-2.5-flash'

# =========================================================
# 2. TEXT EXTRACTION HELPERS
# =========================================================

def extract_text_from_pdf(file_stream):
    """Robust PDF text extraction"""
    try:
        reader = PyPDF2.PdfReader(file_stream)
        text = ""
        # Read all pages. Gemini 1.5 Flash has a 1M token window,
        # so we don't need to limit pages artificially.
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
        return text
    except Exception as e:
        print(f"PDF Extraction Error: {e}")
        return ""

def extract_text_from_docx(file_stream):
    """Robust DOCX text extraction"""
    try:
        doc = docx.Document(file_stream)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text
    except Exception as e:
        print(f"DOCX Extraction Error: {e}")
        return ""

# =========================================================
# 3. DETERMINISTIC PROMPT LOGIC
# =========================================================

def get_consistent_esg_prompt(report_text):
    return f"""
    You are an expert ESG Auditor analyzing a BRSR (Business Responsibility and Sustainability Report).

    ### INSTRUCTIONS FOR CONSISTENCY
    To ensure consistent scoring, you must STRICTLY follow this rubric. Do not guess numbers between buckets.

    ### SCORING RUBRIC (Use ONLY these numbers: 0, 25, 50, 75, 100)
    *   **0**: No mention or "Not Applicable" for critical KPIs.
    *   **25**: Qualitative statements only (e.g., "We aim to save water" but no numbers).
    *   **50**: Policy exists + Basic Quantitative data (e.g., "We used 500kl water").
    *   **75**: Quantitative Data + Year-over-Year comparison showing improvement.
    *   **100**: Quantitative Data + Future Targets (e.g., "Net Zero by 2040") + External Assurance.

    ### ANALYSIS CRITERIA

    **1. ENVIRONMENT (E)**
    - Look for: Scope 1 & 2 Emissions, Renewable Energy %, Zero Liquid Discharge (ZLD), Plastic Waste/EPR.

    **2. SOCIAL (S)**
    - Look for: Gender Diversity (% Women), Safety (LTIFR/Accident rates), CSR Spending amounts.

    **3. GOVERNANCE (G)**
    - Look for: Board Independence, Data Privacy Policies, Anti-Corruption/Whistleblower cases.

    ### OUTPUT FORMAT (JSON ONLY)
    {{
        "company_name": "Name of Company",
        "sector": "Sector (e.g. IT, Manufacturing, Finance)",
        "esg_score": <Calculate Weighted Average: (E*0.4) + (S*0.3) + (G*0.3)>,
        "breakdown": {{
            "E": <Score from 0,25,50,75,100>,
            "S": <Score from 0,25,50,75,100>,
            "G": <Score from 0,25,50,75,100>
        }},
        "reasoning": {{
            "E": "One sentence explaining why you gave this score.",
            "S": "One sentence explaining why you gave this score.",
            "G": "One sentence explaining why you gave this score."
        }},
        "confidence": "High"
    }}

    ### REPORT TEXT
    {report_text[:400000]}
    """

# =========================================================
# 4. API ENDPOINT
# =========================================================

@app.route('/analyze-file', methods=['POST'])
def analyze_file():
    # 1. Validation
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files['file']
    filename = file.filename.lower()

    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    print(f"Processing File: {filename}")

    try:
        # 2. Read & Extract
        file_bytes = file.read()
        file_stream = io.BytesIO(file_bytes)

        full_text = ""
        if filename.endswith('.pdf'):
            full_text = extract_text_from_pdf(file_stream)
        elif filename.endswith('.docx') or filename.endswith('.doc'):
            full_text = extract_text_from_docx(file_stream)
        else:
            return jsonify({"error": "Unsupported format. Send PDF or DOCX."}), 400

        if len(full_text.strip()) < 100:
            return jsonify({
                "esg_score": 0,
                "confidence": "Invalid",
                "analysis_summary": "File empty or unreadable (Scanned PDF?)."
            }), 200

        # 3. Generate Prompt
        prompt = get_consistent_esg_prompt(full_text)
        print("Sending to Gemini...")

        # 4. Call API with Retry Logic & Deterministic Config
        max_retries = 3
        response_text = ""

        for attempt in range(max_retries):
            try:
                response = client.models.generate_content(
                    model=MODEL_NAME,
                    contents=prompt,
                    config=types.GenerateContentConfig(
                        # CRITICAL SETTINGS FOR CONSISTENCY
                        temperature=0.0,  # Zero creativity = Consistent math
                        top_p=1.0,        # Consider all logical tokens
                        top_k=1,          # Only pick the single most likely token
                        max_output_tokens=8192,
                        response_mime_type="application/json"
                    )
                )
                response_text = response.text
                break # Success
            except Exception as e:
                error_str = str(e)
                if "429" in error_str or "RESOURCE_EXHAUSTED" in error_str:
                    wait_time = (attempt + 1) * 5
                    print(f"Rate limit hit. Retrying in {wait_time}s...")
                    time.sleep(wait_time)
                else:
                    raise e # Throw other errors immediately

        # 5. Clean & Return
        raw_json = response_text.strip()

        # Markdown cleanup
        if raw_json.startswith("```json"): raw_json = raw_json[7:]
        if raw_json.startswith("```"): raw_json = raw_json[3:]
        if raw_json.endswith("```"): raw_json = raw_json[:-3]

        result = json.loads(raw_json)
        print("Analysis Complete.")
        return jsonify(result)

    except Exception as e:
        print(f"Error: {e}")
        return jsonify({"error": f"Server Error: {str(e)}"}), 500

if __name__ == '__main__':
    # Threaded=True helps handle multiple requests during testing
    app.run(host='0.0.0.0', port=5000, debug=True, threaded=True)